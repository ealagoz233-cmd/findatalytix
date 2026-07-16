"""Rapor .docx iki dilli: lang=en İngilizce etiket, lang yok/tr Türkçe (geri uyum)."""

import io

from docx import Document

import report


PAYLOAD = {
    "prompt": "THYAO riski nedir?",
    "metrics": {"THYAO.IS": {"model": "GBM", "cagr": 33.22, "vol": 41.5,
                             "sharpe": 0.8, "mdd": -25.3}},
    "dataSources": {"THYAO.IS": "live"},
    "aiText": "Örnek yorum.",
    "aiMeta": {"mode": "live-ai", "analyst": "groq", "referee": "groq",
               "confidence": 80, "tokensIn": 100, "tokensOut": 200},
}


def _all_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_dicts_have_same_schema():
    """TR ve EN sözlükleri birebir aynı anahtarları taşımalı (eksik = KeyError riski)."""
    assert set(report.L["tr"]) == set(report.L["en"])
    assert set(report.L["tr"]["metrics"]) == set(report.L["en"]["metrics"])
    assert set(report.L["tr"]["sources"]) == set(report.L["en"]["sources"])


def test_default_is_turkish():
    """lang gönderilmezse eski davranış: Türkçe rapor + Türkçe dosya adı."""
    data, filename = report.build_report(PAYLOAD)
    text = _all_text(data)
    assert filename.startswith("risk-raporu-")
    assert "Yönetici Özeti" in text
    assert "Sharpe Oranı" in text
    assert "%33,22" in text           # TR sayı biçimi: yüzde önde, virgül
    assert "Canlı Yahoo Finance verisi" in text


def test_english_report():
    """lang=en: başlıklar/etiketler İngilizce, sayı biçimi 33.22%, dosya adı EN."""
    data, filename = report.build_report({**PAYLOAD, "lang": "en"})
    text = _all_text(data)
    assert filename.startswith("risk-report-")
    assert "Executive Summary" in text
    assert "Sharpe Ratio" in text
    assert "33.22%" in text           # EN sayı biçimi: nokta, yüzde sonda
    assert "Live Yahoo Finance data" in text
    assert "Confidence score" in text or "confidence score" in text
    # TR etiket sızıntısı yok
    assert "Yönetici Özeti" not in text
    assert "Hakem" not in text
