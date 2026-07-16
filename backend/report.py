"""
FinDatalytix — report.py (Ay 5: gerçek .docx risk raporu)
=========================================================
Frontend, son simülasyonun tüm durumunu (metrikler, AI yorumu,
hakem skoru, kaynaklar) POST /api/report gövdesinde gönderir;
bu modül profesyonel bir Word raporu üretip bayt olarak döner.
Sunucu durumsuz (stateless) kalır — rapor, o an ekranda görünen
verinin birebir resmi belgesidir.

İki dilli: payload'daki lang ("tr" | "en") arayüz diliyle senkron
gelir; başlıklar, etiketler, sayı/tarih biçimi ve dosya adı ona uyar.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Kurumsal renkler (arayüz paletiyle uyumlu)
NAVY = RGBColor(0x11, 0x1C, 0x31)
GOLD = RGBColor(0x8A, 0x6D, 0x2F)
GREY = RGBColor(0x5A, 0x64, 0x72)
RED = RGBColor(0xB0, 0x3A, 0x3A)

METRIC_ROWS = [
    ("model", None),
    ("cagr", "%"),
    ("vol", "%"),
    ("sharpe", ""),
    ("mdd", "%"),
]

# Rapor sözlüğü — arayüzdeki FDX.I18N ile aynı ilke: her görünen metin
# iki dilde, kod tek. Anahtar eksiği testte yakalanır (iki dil aynı şema).
L = {
    "tr": {
        "filename": "risk-raporu",
        "title": "FinDatalytix — Karşılaştırmalı Risk Raporu",
        "created": "Oluşturulma",
        "engine": "Motor: Monte Carlo (GBM, 2.000 yol)",
        "question": "Analiz sorusu",
        "h1": "1. Yönetici Özeti",
        "noAi": "Bu rapor oluşturulduğunda AI yorumu mevcut değildi; "
                "aşağıdaki sayısal bulgular esas alınmalıdır.",
        "h2": "2. Metrik Karşılaştırması",
        "metricCol": "Metrik",
        "winner": "Risk-ayarlı getiri (Sharpe) bazında öne çıkan: {sym}.",
        "h3": "3. Veri Kaynakları",
        "assetRow": "Varlık {key}: {src}",
        "noSrc": "Veri kaynağı bilgisi iletilmedi.",
        "h4": "4. Yapay Zeka Değerlendirmesi",
        "analyst": "Analist model: {m}",
        "refScore": "Hakem ({r}) güven skoru: {c}/100",
        "refNote": "Hakem notu: {n}",
        "tokens": "Token kullanımı: {i} giriş / {o} çıkış",
        "errFallback": "AI servisine bu çalıştırmada ulaşılamadı (kota/ağ hatası); "
                       "yorum ham sayısal sonuçlarla sınırlıdır. Sayısal bulgular geçerlidir.",
        "template": "Bu raporun yorumu şablon modunda üretildi (AI anahtarı tanımlı değildi). "
                    "Gerçek analist + hakem değerlendirmesi için .env dosyasına API anahtarı ekleyin.",
        "h5": "5. Kullanılan Doküman Kaynakları (RAG)",
        "noRag": "Bu analizde indeksli doküman kullanılmadı.",
        "legal": "Bu rapor FinDatalytix tarafından otomatik üretilmiştir; simülasyon sonuçları "
                 "istatistiksel modellere dayanır, geleceği garanti etmez ve yatırım tavsiyesi değildir.",
        "dateFmt": "%d.%m.%Y %H:%M",
        "metrics": {
            "model": "Model",
            "cagr": "Yıllık Getiri (CAGR)",
            "vol": "Volatilite (σ)",
            "sharpe": "Sharpe Oranı",
            "mdd": "Maksimum Düşüş (MDD)",
        },
        "sources": {
            "live": "Canlı Yahoo Finance verisi",
            "cache": "Önbellekteki piyasa verisi (≤1 saat)",
            "fallback": "Varsayılan parametreler (canlı veri alınamadı)",
        },
    },
    "en": {
        "filename": "risk-report",
        "title": "FinDatalytix — Comparative Risk Report",
        "created": "Generated",
        "engine": "Engine: Monte Carlo (GBM, 2,000 paths)",
        "question": "Analysis question",
        "h1": "1. Executive Summary",
        "noAi": "No AI commentary was available when this report was generated; "
                "rely on the numerical findings below.",
        "h2": "2. Metric Comparison",
        "metricCol": "Metric",
        "winner": "Best risk-adjusted return (Sharpe): {sym}.",
        "h3": "3. Data Sources",
        "assetRow": "Asset {key}: {src}",
        "noSrc": "No data-source information was provided.",
        "h4": "4. AI Assessment",
        "analyst": "Analyst model: {m}",
        "refScore": "Referee ({r}) confidence score: {c}/100",
        "refNote": "Referee note: {n}",
        "tokens": "Token usage: {i} in / {o} out",
        "errFallback": "The AI service could not be reached during this run (quota/network error); "
                       "commentary is limited to raw numerical results, which remain valid.",
        "template": "This report's commentary was produced in template mode (no AI key configured). "
                    "Add an API key to the .env file for real analyst + referee assessment.",
        "h5": "5. Document Sources Used (RAG)",
        "noRag": "No indexed documents were used in this analysis.",
        "legal": "This report was generated automatically by FinDatalytix; simulation results are "
                 "based on statistical models, do not guarantee future outcomes and do not "
                 "constitute investment advice.",
        "dateFmt": "%Y-%m-%d %H:%M",
        "metrics": {
            "model": "Model",
            "cagr": "Annual Return (CAGR)",
            "vol": "Volatility (σ)",
            "sharpe": "Sharpe Ratio",
            "mdd": "Maximum Drawdown (MDD)",
        },
        "sources": {
            "live": "Live Yahoo Finance data",
            "cache": "Cached market data (≤1 hour)",
            "fallback": "Default parameters (live data unavailable)",
        },
    },
}


def _fmt(value, suffix, lang="tr"):
    if suffix is None:            # metin alanı (model adı)
        return str(value)
    if isinstance(value, (int, float)):
        text = f"{value:.2f}"
        if lang == "tr":
            # Türkçe ondalık ve yüzde işareti önde: 33.22 -> "%33,22"
            text = text.replace(".", ",")
            return f"%{text}" if suffix == "%" else text
        return f"{text}%" if suffix == "%" else text
    return str(value)


def _heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def _para(doc: Document, text: str, size=11, color=None, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def build_report(payload: dict) -> tuple[bytes, str]:
    """(docx_baytları, dosya_adı) döner."""

    lang = "en" if payload.get("lang") == "en" else "tr"
    t = L[lang]

    metrics = payload.get("metrics") or {}
    asset_keys = list(metrics.keys())          # dinamik: 1..N varlık
    ai_text = (payload.get("aiText") or "").strip()
    meta = payload.get("aiMeta") or {}
    prompt = (payload.get("prompt") or "").strip()
    data_sources = payload.get("dataSources") or {}

    now = datetime.now()
    filename = f"{t['filename']}-{now:%Y-%m-%d-%H%M}.docx"

    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Kapak başlığı ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(t["title"])
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp = now.strftime(t["dateFmt"])
    run = sub.add_run(f"{t['created']}: {stamp}  ·  {t['engine']}")
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    if prompt:
        _para(doc, f"{t['question']}: “{prompt}”", size=10, color=GREY, italic=True)

    # ---- 1. Yönetici Özeti ----
    _heading(doc, t["h1"])
    if ai_text:
        _para(doc, ai_text)
    else:
        _para(doc, t["noAi"], color=GREY, italic=True)

    # ---- 2. Metrik Karşılaştırması (dinamik sütun) ----
    _heading(doc, t["h2"])
    table = doc.add_table(rows=1, cols=1 + max(len(asset_keys), 1))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate([t["metricCol"]] + (asset_keys or ["—"])):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(text)
        run.bold = True

    for key, suffix in METRIC_ROWS:
        if key == "model":
            continue   # sütun başlığı zaten sembolün kendisi
        row = table.add_row().cells
        row[0].text = t["metrics"][key]
        for i, sym in enumerate(asset_keys):
            value = metrics[sym].get(key, "—")
            cell = row[i + 1]
            cell.text = ""
            run = cell.paragraphs[0].add_run(_fmt(value, suffix, lang))
            if key == "mdd" and isinstance(value, (int, float)) and value < 0:
                run.font.color.rgb = RED

    # Risk-ayarlı kazanan
    try:
        if asset_keys:
            winner = max(asset_keys, key=lambda k: float(metrics[k].get("sharpe", 0)))
            _para(doc, t["winner"].format(sym=winner), bold=True, color=GOLD)
    except (TypeError, ValueError):
        pass

    # ---- 3. Veri Kaynakları ----
    _heading(doc, t["h3"])
    if data_sources:
        for asset_key, src in data_sources.items():
            doc.add_paragraph(
                t["assetRow"].format(key=asset_key, src=t["sources"].get(src, src)),
                style="List Bullet",
            )
    else:
        _para(doc, t["noSrc"], color=GREY, italic=True)

    # ---- 4. Yapay Zeka Değerlendirmesi ----
    _heading(doc, t["h4"])
    mode = meta.get("mode")
    if mode == "live-ai":
        doc.add_paragraph(t["analyst"].format(m=meta.get("analyst", "—")),
                          style="List Bullet")
        conf = meta.get("confidence")
        referee = meta.get("referee", "—")
        if conf is not None:
            doc.add_paragraph(t["refScore"].format(r=referee, c=conf),
                              style="List Bullet")
        if meta.get("refereeNote"):
            doc.add_paragraph(t["refNote"].format(n=meta["refereeNote"]),
                              style="List Bullet")
        doc.add_paragraph(
            t["tokens"].format(i=meta.get("tokensIn", 0), o=meta.get("tokensOut", 0)),
            style="List Bullet",
        )
    elif mode == "error-fallback":
        # Dürüstlük: anahtar VARDI ama çağrı başarısız oldu — "anahtar yok"
        # demek yanlış teşhis koydurur (kullanıcı boşuna .env kurcalar).
        _para(doc, t["errFallback"], color=GREY, italic=True)
    else:
        _para(doc, t["template"], color=GREY, italic=True)

    # ---- 5. RAG Kaynakları ----
    rag_sources = meta.get("ragSources") or []
    _heading(doc, t["h5"])
    if rag_sources:
        for src in rag_sources:
            doc.add_paragraph(src, style="List Bullet")
    else:
        _para(doc, t["noRag"], color=GREY, italic=True)

    # ---- Yasal uyarı ----
    doc.add_paragraph()
    _para(doc, t["legal"], size=8.5, color=GREY, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), filename
